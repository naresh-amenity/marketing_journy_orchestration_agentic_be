import json
import os
from datetime import datetime
from typing import Any, Dict, List, Optional

import requests
from docx import Document
from docx.shared import Inches, Pt

from app.models.model import ToolResponse
from app.tools.base_tool import BaseTool


class JourneyTool(BaseTool):
    """
    Tool for creating and managing marketing journeys

    Provides functionality to:
    - Create new marketing journeys
    - Check the status of existing journeys
    - Get journey reports and generate documents
    - Process journey reports with persona files
    """

    def __init__(self):
        """
        Initialize the Journey Tool with the API client
        """
        self.base_url = "http://localhost:8003/api/boostt_ai"
        self.output_dir = "journey_reports"

        # Create output directory if it doesn't exist
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def get_name(self) -> str:
        """
        Get the name of the tool

        Returns:
            The name of the tool
        """
        return "journy_tool"

    def get_description(self) -> str:
        """
        Get the description of the tool

        Returns:
            The description of the tool
        """
        return """
        Tool for creating and managing marketing journeys.
        Can create new journeys, check status, generate reports, and update journeys with persona files.
        
        Actions:
        - create_journey: Create a new marketing journey
        - get_journey_status: Check status of an existing journey
        - get_journey_report: Get a journey report and generate a document
        - check_journey_report: Process a journey report with persona files
        """

    def get_required_params(self) -> List[str]:
        """
        Get the list of required parameters for the tool

        Returns:
            A list of required parameter names
        """
        return ["action"]

    def get_optional_params(self) -> List[str]:
        """
        Get the list of optional parameters for the tool

        Returns:
            A list of optional parameter names
        """
        return [
            "user_problem_statement",
            "budget",
            "target_id",
            "session_token",
            "name",
            "date",
            "journey_id",
            "mongo_id",
            "persona_files",
            "uploaded_document",
            "base_url",
            "document_action",
            "document_name",
            "document_content",
            "request_download",
            "has_audience_model",
        ]

    async def execute(self, parameters: Dict[str, Any]) -> ToolResponse:
        """
        Execute the tool with the given parameters

        Args:
            parameters: The parameters for the tool

        Returns:
            The response from the tool
        """
        # Check required parameters
        missing_params = self.check_required_params(parameters)
        if missing_params:
            return ToolResponse(
                status="error",
                message="I'd be happy to help with your marketing journey. What would you like to do? You can create a new journey, check status of an existing journey, get a journey report, or update a journey report with persona files.",
                required_inputs=["action"],
                needed_input=["action"],
            )

        action = parameters.get("action")
        print("actionaction, action", parameters)

        try:
            # Handle document operations first
            if action == "document_operation":
                document_action = parameters.get("document_action")

                if document_action == "download":
                    return await self._download_journey_document(parameters)
                elif document_action == "upload":
                    return await self._upload_journey_document(parameters)
                else:
                    return ToolResponse(
                        status="error",
                        message=f"Unknown document action: {document_action}. Supported actions are 'download' and 'upload'.",
                        required_inputs=["document_action"],
                        needed_input=["document_action"],
                    )

            # Handle regular journey operations
            if action == "create_journey":
                return await self._create_journey(parameters)
            elif action == "get_journey_status":
                return await self._get_journey_status(parameters)
            elif action == "get_journey_report":
                return await self._get_journey_report(parameters)
            elif action == "check_journey_report":
                return await self._check_journey_report(parameters)
            else:
                return ToolResponse(
                    status="error",
                    message=f"I don't recognize that action. You can choose from: 'create_journey', 'get_journey_status', 'get_journey_report', 'check_journey_report', or 'document_operation'. What would you like to do?",
                    required_inputs=["action"],
                    needed_input=["action"],
                )
        except Exception as e:
            return ToolResponse(
                status="error",
                message=f"I encountered an error while processing your request: {str(e)}. Could you try again?",
                data={"error": str(e)},
            )

    async def _create_journey(self, parameters: Dict[str, Any]) -> ToolResponse:
        """
        Create a new marketing journey

        Args:
            parameters: The parameters for creating a journey

        Returns:
            ToolResponse with the journey creation result
        """
        # Check for problem_statement and map to user_problem_statement if needed
        if (
            "problem_statement" in parameters
            and "user_problem_statement" not in parameters
        ):
            parameters["user_problem_statement"] = parameters["problem_statement"]

        if "has_audience_model" not in parameters:
            return ToolResponse(
                status="error",
                message="Do you have a model for the audience?",
                required_inputs=["has_audience_model"],
                needed_input=["has_audience_model"],
            )

        if "has_audience_model" in parameters:
            if (
                parameters["has_audience_model"] is None
                or parameters["has_audience_model"] == ""
            ):
                return ToolResponse(
                    status="error",
                    message="Do you have a model for the audience?",
                    required_inputs=["has_audience_model"],
                    needed_input=["has_audience_model"],
                )

        # Define required parameters based on whether user has a model
        required_params = [
            "user_problem_statement",
            "budget",
            "target_id",
            "session_token",
            "name",
            "date",
        ]

        # If user has a model, add model_id to required parameters
        has_model = str(parameters.get("has_audience_model", "")).lower()
        if has_model == "yes":
            required_params.append("model_id")

        print("Required parameters:", required_params)
        print("Has audience model:", has_model)

        missing = [
            param
            for param in required_params
            if param not in parameters or not parameters[param]
        ]

        if missing:
            # Group related parameters for a more conversational flow
            if "model_id" in missing and "model_id" in required_params:
                return ToolResponse(
                    status="error",
                    message="To create a new marketing journey, I'll need some information. First, could you please provide the model for this journey?",
                    required_inputs=["model_id"],
                    needed_input=["model_id"],
                )
            elif "target_id" in missing:
                return ToolResponse(
                    status="error",
                    message=f"I'll need some information. First, could you please provide the audience for this journey?",
                    required_inputs=["target_id"],
                    needed_input=["target_id"],
                )
            elif "user_problem_statement" in missing:
                return ToolResponse(
                    status="error",
                    message=f"I'll need your problem statement. could you please provide the problem statement for this journey?",
                    required_inputs=["user_problem_statement"],
                    needed_input=["user_problem_statement"],
                )
            elif "budget" in missing or "name" in missing:
                tech_params = [
                    param for param in ["budget", "name"] if param in missing
                ]
                return ToolResponse(
                    status="error",
                    message=f"I'll need some information. Could you please provide {', '.join(tech_params)} for this journey?",
                    required_inputs=tech_params,
                    needed_input=tech_params,
                )
            # elif "budget" in missing:
            #     return ToolResponse(
            #         status="error",
            #         message=f"Thanks for providing the problem statement. What's the total budget allocated for this marketing journey?",
            #         required_inputs=["budget"],
            #         needed_input=["budget"]
            #     )
            # elif "name" in missing:
            #     return ToolResponse(
            #         status="error",
            #         message="Great! Now, what would you like to name this marketing journey?",
            #         required_inputs=["name"],
            #         needed_input=["name"]
            #     )
            elif "date" in missing:
                return ToolResponse(
                    status="error",
                    message="When would you like this journey to start? Please select a date",
                    required_inputs=["date"],
                    needed_input=["date"],
                )
            elif "session_token" in missing:
                # Technical parameters that can be requested together

                return ToolResponse(
                    status="error",
                    message=f"Almost there! To complete the journey setup, I need session token",
                    required_inputs=["session_token"],
                    needed_input=["session_token"],
                )

        try:
            endpoint = f"{self.base_url}/create_journey/"

            payload = {
                "user_problem_statement": parameters["user_problem_statement"],
                "budget": int(parameters["budget"]),
                "TARGET_ID": parameters["target_id"],
                "SESSION_TOKEN": parameters["session_token"],
                "name": parameters["name"],
                "date": parameters["date"],
                "user_id": parameters["user_id"],
                "model_id_persona": parameters["model_id"],
                "model_id_propensity": parameters["model_id"],
            }

            response = requests.post(endpoint, json=payload)
            response.raise_for_status()
            result = response.json()
            # result = {"Journey_process_id": "6808b42aec2d512c61e8d848", "journey_status": "not_started", "journeyID": "7rc18tk2tc"}

            return ToolResponse(
                status="success",
                message="🛠️✨ We're building your journey! Hang tight your planning will be ready in just a few minutes.",
                data={
                    "journey_process_id": result.get("Journey_process_id"),
                    "journey_id": result.get("journeyID"),
                    "status": result.get("status"),
                },
            )
        except Exception as e:
            return ToolResponse(
                status="error",
                message=f"Error creating journey: {str(e)}",
                data={"error": str(e)},
            )

    async def _get_journey_status(self, parameters: Dict[str, Any]) -> ToolResponse:
        """
        Get the status of a journey

        Args:
            parameters: The parameters for getting a journey status

        Returns:
            ToolResponse with the journey status
        """
        if "mongo_id" not in parameters or not parameters["mongo_id"]:
            return ToolResponse(
                status="error",
                message="No journey created yet would you like to create a new journey?",
            )

        try:
            endpoint = f"{self.base_url}/status/{parameters['mongo_id']}"

            response = requests.get(endpoint)
            response.raise_for_status()
            result = response.json()

            # Create a descriptive message based on the journey status information
            journey_status = result.get("journey_status", "Unknown")
            journey_name = result.get("name", "Your marketing journey")
            journey_date = result.get("date", "N/A")
            journey_id = result.get("journey_id", "Unknown")

            # message = f"Journey '{journey_name}'"

            # # Only add date if it's available and valid
            # if journey_date and journey_date != "N/A":
            #     message += f" (starting {journey_date})"

            # message += ":\n\n"

            # Add specific details based on status
            if (
                journey_status.lower() == "not_started"
                or journey_status.lower() == "processing"
            ):
                message = f"🌍✨ Your journey is taking shape! Please check back in a few minutes"
            elif journey_status.lower() == "completed":
                message = (
                    f"✅ Journey Completed! You can check it from the journey page"
                )
            elif journey_status.lower() == "user_approved":
                message = f"🛠️📧 We’re configuring your journey! Please check back in a few hours we’ll send you an email as soon as everything’s ready to go."
            elif journey_status.lower() == "document created":
                message = f"📄✅ Your journey report is ready! Please download the report and review the planning. Once you're happy with it, go ahead and approve!"
            elif journey_status.lower() == "error":
                message = f"❌ Status: Failed Uh-oh! Something went wrong while processing your journey. Please create a new journey"
            else:
                message = f"Status: {journey_status}"

            return ToolResponse(
                status="success",
                message=message,
                data=None,  # No data included as requested
            )
        except Exception as e:
            return ToolResponse(
                status="error",
                message=f"Error getting journey status: {str(e)}",
                data={"error": str(e)},
            )

    async def _get_journey_report(self, parameters: Dict[str, Any]) -> ToolResponse:
        """
        Get a journey report

        Args:
            parameters: The parameters for getting a journey report

        Returns:
            ToolResponse with the journey report and document content for direct download
        """
        if "journey_id" not in parameters or not parameters["journey_id"]:
            return ToolResponse(
                status="error",
                message="No journey created yet would you like to create a new journey?",
            )

        try:
            endpoint = f"{self.base_url}/journey_report/"

            params = {"journeyID": parameters["journey_id"]}

            response = requests.get(endpoint, params=params)
            response.raise_for_status()
            result = response.json()

            # Generate Word document from the journey report
            doc_metadata = self._generate_journey_doc(result, parameters["journey_id"])

            # Always include document content for direct download by frontend
            # Read the document as binary data
            with open(doc_metadata["file_path"], "rb") as file:
                binary_data = file.read()

            # Convert binary data to base64
            import base64

            document_content = base64.b64encode(binary_data).decode("utf-8")

            return ToolResponse(
                status="success",
                message=f"Journey report retrieved successfully.",
                data={
                    "journey_id": parameters["journey_id"],
                    "report_data": result,
                    "document": doc_metadata,
                    "document_content": document_content,
                    "filename": doc_metadata["file_name"],
                    "content_type": doc_metadata["content_type"],
                    "file_size": doc_metadata["file_size"],
                },
            )
        except Exception as e:
            return ToolResponse(
                status="error",
                message=f"Error getting journey report: {str(e)}",
                data={"error": str(e)},
            )

    async def _check_journey_report(self, parameters: Dict[str, Any]) -> ToolResponse:
        """
        Check and update journey report with persona files or an uploaded document

        Args:
            parameters: The parameters for checking and updating a journey report

        Returns:
            ToolResponse with the result of the check
        """
        if "journey_id" not in parameters or not parameters["journey_id"]:
            return ToolResponse(
                status="error",
                message="No journey created yet would you like to create a new journey?",
            )

        try:
            # if True:
            endpoint = f"{self.base_url}/check_journey_report/"

            data = {"journeyID": parameters["journey_id"]}
            files = []

            # Handle uploaded document if provided
            if "uploaded_document" in parameters and parameters["uploaded_document"]:
                doc_path = parameters["uploaded_document"]

                # Check if file exists
                if not os.path.exists(doc_path):
                    return ToolResponse(
                        status="error",
                        message=f"I couldn't find the uploaded document at {doc_path}. Please ensure the file path is correct.",
                        data={"error": f"File not found: {doc_path}"},
                    )

                # Extract text content directly
                doc_text = self._extract_text_from_docx(doc_path)

                # Send text content instead of file
                files.append(("document", ("document.txt", doc_text, "text/plain")))

                # Set document_provided flag in data
                data["document_provided"] = "true"
            # Handle document_content if provided (for uploads via process endpoint)
            elif "document_content" in parameters and parameters["document_content"]:
                # Create a temporary file
                import base64
                import tempfile

                timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = (
                    f"journey_{parameters['journey_id']}_{timestamp_str}_modified.docx"
                )
                file_path = os.path.join(self.output_dir, filename)

                # Decode base64 content and save to file
                binary_data = base64.b64decode(parameters["document_content"])
                with open(file_path, "wb") as f:
                    f.write(binary_data)

                # Extract text content directly
                doc_text = self._extract_text_from_docx(file_path)

                # Send text content instead of file
                files.append(("document", ("document.txt", doc_text, "text/plain")))

                # Set document_provided flag in data
                data["document_provided"] = "true"

                # Clean up temporary file
                try:
                    os.remove(file_path)
                except:
                    pass
            # else:
            #     # No document provided, use existing document and persona files if any
            #     data["document_provided"] = "false"

            #     # Handle persona files if provided
            #     if "persona_files" in parameters and parameters["persona_files"]:
            #         for i, file_path in enumerate(parameters["persona_files"]):
            #             if os.path.exists(file_path):
            #                 with open(file_path, 'rb') as f:
            #                     files.append(('files', (f'persona_{i}.txt', f, 'text/plain')))
            #             else:
            #                 return ToolResponse(
            #                     status="error",
            #                     message=f"I couldn't find one of the persona files you specified: {file_path}. Could you please check the file path and try again?",
            #                     data={"error": f"File not found: {file_path}"}
            #                 )
            print("data **************88", data)
            print("files **************88", files)
            response = requests.post(endpoint, data=data, files=files)
            response.raise_for_status()
            result = response.json()

            # Generate appropriate message based on whether a document was uploaded
            if "uploaded_document" in parameters and parameters["uploaded_document"]:
                message = (
                    "Your modified journey document has been processed successfully."
                )
            elif "document_content" in parameters and parameters["document_content"]:
                message = (
                    "Your modified journey document has been processed successfully."
                )
            else:
                message = (
                    "Journey report processed successfully with the existing document."
                )

            return ToolResponse(
                status="success",
                message=message,
                data={
                    "journey_id": parameters["journey_id"],
                    "status": result.get("status", "processed"),
                    "message": result.get("message", "Report updated successfully"),
                },
            )
        except Exception as e:
            return ToolResponse(
                status="error",
                message=f"Error processing journey report: {str(e)}",
                data={"error": str(e)},
            )

    def _generate_journey_doc(
        self, report_data: Dict[str, Any], journey_id: str
    ) -> Dict[str, Any]:
        """
        Generate a Word document from the journey report data.

        Args:
            report_data: The journey report data from the API
            journey_id: The journey ID for naming the file

        Returns:
            Dictionary containing document metadata and file path
        """
        # Create a new Document
        doc = Document()

        # Add document title with styling
        title = doc.add_heading(f"Marketing Journey Report", 0)

        # Add timestamp
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        timestamp_para = doc.add_paragraph()
        timestamp_para.add_run(f"Generated on: {current_time}").italic = True

        # Add journey ID section
        doc.add_heading("Journey Information", level=1)
        doc.add_paragraph(f"Journey ID: {journey_id}")

        # Format the document based on the specific structure
        if "journey_report" in report_data:
            journey_report = report_data["journey_report"]

            # Process non-persona data if available
            if "non_persona" in journey_report:
                non_persona = journey_report["non_persona"]

                # Add budget information if available
                if "budget" in non_persona:
                    budget_heading = doc.add_heading("Budget", level=1)
                    doc.add_paragraph(f'Total Budget: ${non_persona["budget"]:,}')

                # Process user content which contains the markdown-formatted journey map
                if "user_content" in non_persona:
                    user_content = non_persona["user_content"]

                    # Parse and format the markdown content
                    self._format_markdown_content(doc, user_content)
        # Fall back to checking the expected structure from previous implementation
        elif "user_report" in report_data:
            user_report = report_data["user_report"]

            # Add simple paragraph with raw data
            doc.add_paragraph(str(user_report))
        else:
            # Add raw data if structured data is not available
            doc.add_heading("Report Data", level=1)
            doc.add_paragraph(str(report_data))

        # Generate filename with timestamp to avoid overwriting
        timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_name = f"journey_{journey_id}_{timestamp_str}.docx"
        file_path = os.path.join(self.output_dir, file_name)

        # Save the document
        doc.save(file_path)

        # Get file size for metadata
        file_size = os.path.getsize(file_path)

        # Return metadata dictionary
        return {
            "file_path": file_path,
            "file_name": file_name,
            "journey_id": journey_id,
            "timestamp": timestamp_str,
            "file_size": file_size,
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }

    def _format_markdown_content(self, doc: Document, content: str) -> None:
        """
        Format markdown content into a Word document with proper styling.

        Args:
            doc: The document to add content to
            content: Markdown formatted content
        """
        # Split content by lines
        lines = content.strip().split("\n")

        for line in lines:
            # Skip empty lines
            if not line.strip():
                continue

            # Handle headings with different levels
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                # Remove any ** markdown from the heading
                heading_text = line[3:].replace("**", "")
                doc.add_heading(heading_text, level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("#### "):
                doc.add_heading(line[5:], level=4)
            # Handle bold step headings (like "**Step 1: Something**")
            elif line.strip().startswith("**Step ") and "**" in line:
                # Extract just the step text without markdown
                step_text = line.strip().replace("**", "")
                p = doc.add_heading(step_text, level=3)
            # Handle other bold text as subheadings
            elif line.startswith("**") and line.endswith("**") and "**:" not in line:
                p = doc.add_paragraph()
                bold_text = line.strip("*")
                p.add_run(bold_text).bold = True
            # Handle list items
            elif line.strip().startswith("* "):
                # Calculate indentation level
                indent_level = len(line) - len(line.lstrip())
                level = 0

                # Convert spaces to indentation level
                if indent_level > 0:
                    level = indent_level // 2

                # Extract the text content of the list item
                text = line.strip().lstrip("* ")

                # Create bullet with proper indentation using Word's built-in bullet style
                p = doc.add_paragraph(style="List Bullet")

                # Set proper indentation
                if level == 0:
                    # Top level bullets should be at the left margin
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.first_line_indent = Inches(-0.25)
                else:
                    # Nested bullets should be indented more
                    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
                    p.paragraph_format.first_line_indent = Inches(-0.25)

                # Format the text of the bullet point
                if text.startswith("**") and "**:" in text:
                    # Handle "**Key:** Value" format
                    parts = text.split("**:", 1)
                    key = parts[0].strip("* ")
                    value = parts[1].strip() if len(parts) > 1 else ""

                    # Remove any remaining ** characters
                    key = key.replace("**", "")

                    # Add the formatted text
                    p.add_run(f"{key}: ").bold = True
                    p.add_run(value)
                elif "**" in text:
                    # Process text with bold segments
                    segments = text.split("**")
                    is_bold = False

                    for segment in segments:
                        if segment:
                            run = p.add_run(segment)
                            run.bold = is_bold
                        is_bold = not is_bold
                else:
                    # Regular bullet item
                    p.add_run(text)
            else:
                # Regular paragraph text
                doc.add_paragraph(line)

    async def _download_journey_document(
        self, parameters: Dict[str, Any]
    ) -> ToolResponse:
        """
        Download a journey document

        Args:
            parameters: The parameters for downloading a journey document

        Returns:
            ToolResponse with the journey document data
        """
        if "journey_id" not in parameters or not parameters["journey_id"]:
            return ToolResponse(
                status="error",
                message="I'll help you download your journey document. Could you provide the journey ID?",
                required_inputs=["journey_id"],
                needed_input=["journey_id"],
            )

        try:
            journey_id = parameters["journey_id"]
            document_name = parameters.get("document_name")

            # Determine the file path based on journey_id
            if document_name:
                # If a specific document name is provided
                file_path = os.path.join(self.output_dir, document_name)
                if not os.path.exists(file_path):
                    return ToolResponse(
                        status="error",
                        message=f"I couldn't find the document named {document_name}. Please check the name and try again.",
                        data={"error": f"Document not found: {document_name}"},
                    )
            else:
                # Find the latest document for this journey_id
                import glob

                pattern = os.path.join(self.output_dir, f"journey_{journey_id}_*.docx")
                matching_files = glob.glob(pattern)

                if not matching_files:
                    return ToolResponse(
                        status="error",
                        message=f"I couldn't find any documents for journey {journey_id}. Please check the journey ID and try again.",
                        data={"error": f"No documents found for journey {journey_id}"},
                    )

                # Sort by creation time (newest first)
                matching_files.sort(key=lambda x: os.path.getctime(x), reverse=True)
                file_path = matching_files[0]

            # Read the document as binary data
            with open(file_path, "rb") as file:
                binary_data = file.read()

            # Convert binary data to base64
            import base64

            document_content = base64.b64encode(binary_data).decode("utf-8")

            # Get file metadata
            file_size = os.path.getsize(file_path)
            file_name = os.path.basename(file_path)

            return ToolResponse(
                status="success",
                message=f"Journey document retrieved successfully.",
                data={
                    "journey_id": journey_id,
                    "document_name": file_name,
                    "document_content": document_content,
                    "file_size": file_size,
                    "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                },
            )
        except Exception as e:
            return ToolResponse(
                status="error",
                message=f"Error downloading journey document: {str(e)}",
                data={"error": str(e)},
            )

    async def _upload_journey_document(
        self, parameters: Dict[str, Any]
    ) -> ToolResponse:
        """
        Upload a journey document

        Args:
            parameters: The parameters for uploading a journey document

        Returns:
            ToolResponse with the upload result
        """
        if "journey_id" not in parameters or not parameters["journey_id"]:
            return ToolResponse(
                status="error",
                message="I'll help you upload your journey document. Could you provide the journey ID?",
                required_inputs=["journey_id"],
                needed_input=["journey_id"],
            )

        if "document_content" not in parameters or not parameters["document_content"]:
            return ToolResponse(
                status="error",
                message="I need the document content to upload. Please provide the document as base64-encoded content.",
                required_inputs=["document_content"],
                needed_input=["document_content"],
            )

        try:
            journey_id = parameters["journey_id"]

            # Ensure upload directory exists
            os.makedirs(self.output_dir, exist_ok=True)

            # Generate a filename with timestamp to avoid overwriting
            timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"journey_{journey_id}_{timestamp_str}_modified.docx"
            file_path = os.path.join(self.output_dir, filename)

            # Decode base64 content and save to file
            import base64

            binary_data = base64.b64decode(parameters["document_content"])
            with open(file_path, "wb") as f:
                f.write(binary_data)

            # Return success response with file details
            return ToolResponse(
                status="success",
                message="Document uploaded successfully.",
                data={
                    "journey_id": journey_id,
                    "file_name": filename,
                    "file_path": file_path,
                    "file_size": os.path.getsize(file_path),
                    "timestamp": timestamp_str,
                },
            )
        except Exception as e:
            return ToolResponse(
                status="error",
                message=f"Error uploading journey document: {str(e)}",
                data={"error": str(e)},
            )

    def _extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract text content from a Word document

        Args:
            docx_path: Path to the Word document

        Returns:
            String containing the extracted text
        """
        try:
            # Load the document
            doc = Document(docx_path)

            # Extract text content
            content = []
            for para in doc.paragraphs:
                content.append(para.text)

            return "\n".join(content)
        except Exception as e:
            # Return empty string if extraction fails
            print(f"Error extracting text from document: {str(e)}")
            return ""
